# SPDX-License-Identifier: MIT
"""
Multimodal Knowledge Graph - Text Knowledge Graph Generator
Paper-to-Knowledge-Graph conversion using LLM-based Skills
"""
import uuid
import json
import re
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional, Tuple, Union
from pathlib import Path
from datetime import datetime
from enum import Enum
import hashlib

from loguru import logger

from app.core.config import settings
from app.models.database import (
    KGEntity, KGRelationship, KnowledgeGraph,
    get_sync_session, Base, sync_engine
)


# =============================================================================
# Data Classes for Skill Outputs
# =============================================================================

class EntityType(Enum):
    """Types of knowledge graph entities"""
    CONCEPT = "concept"
    THEORY = "theory"
    VARIABLE = "variable"
    PARAMETER = "parameter"
    BOUNDARY = "boundary"
    EXTERNAL_REF = "external_ref"
    CONTRIBUTION = "contribution"


class VariableType(Enum):
    """Types of variables in research"""
    IV = "independent_variable"  # Independent Variable
    DV = "dependent_variable"    # Dependent Variable
    MEDIATOR = "mediator"
    MODERATOR = "moderator"
    CONTROL = "control_variable"
    COVARIATE = "covariate"


class RelationshipLabel(Enum):
    """Types of variable relationships"""
    POSITIVE = "positive"
    NEGATIVE = "negative"
    U_SHAPE = "u_shape"
    INVERTED_U = "inverted_u"
    CURVILINEAR = "curvilinear"
    NO_EFFECT = "no_effect"
    MEDIATED = "mediated"
    MODERATED = "moderated"


class ClusterType(Enum):
    """Knowledge graph cluster types"""
    CONTEXT = "context"
    THEORY = "theory"
    METHODOLOGY = "methodology"
    FINDINGS = "findings"
    CONTRIBUTION = "contribution"


class HypothesisStatus(Enum):
    """Status of hypothesis testing"""
    SUPPORTED = "supported"
    PARTIAL = "partial"
    REJECTED = "rejected"


@dataclass
class KGEntityNode:
    """Knowledge Graph Entity from Skill 1"""
    id: str
    name: str
    entity_type: str
    description: Optional[str] = None
    source_section: Optional[str] = None
    confidence: float = 1.0
    properties: Dict[str, Any] = field(default_factory=dict)


@dataclass
class TheoryNode:
    """Theory/Framework node from Skill 2"""
    id: str
    name: str
    theory_type: str  # 'parent_theory', 'framework', 'model'
    parent_theory_id: Optional[str] = None
    description: Optional[str] = None
    source_citation: Optional[str] = None
    confidence: float = 1.0
    properties: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Variable:
    """Research variable from Skill 3"""
    id: str
    name: str
    variable_type: VariableType
    description: Optional[str] = None
    measurement: Optional[str] = None
    relationships: List[Dict] = field(default_factory=list)  # [{'target': str, 'type': RelationshipLabel, 'label': str}]
    confidence: float = 1.0


@dataclass
class ClusterAssignment:
    """Cluster assignment from Skill 4"""
    clusters: Dict[str, List[str]]  # {cluster_type: [entity_ids]}
    unassigned: List[str] = field(default_factory=list)
    cluster_labels: Dict[str, str] = field(default_factory=dict)


@dataclass
class Parameter:
    """Research parameter from Skill 5"""
    id: str
    name: str
    param_type: str  # 'sample_size', 'location', 'tool', 'method', 'time_period'
    value: Any
    unit: Optional[str] = None
    description: Optional[str] = None
    confidence: float = 1.0


@dataclass
class EvidenceMapping:
    """Evidence mapping from Skill 6"""
    variable_relationships: List[Dict]  # [{'source': str, 'target': str, 'status': HypothesisStatus, 'p_value': float, 'beta': float, 'evidence': str}]
    overall_finding: str


@dataclass
class Boundary:
    """Boundary condition from Skill 7"""
    id: str
    boundary_type: str  # 'temporal', 'spatial', 'demographic', 'contextual'
    description: str
    scope: Optional[str] = None  # 'generalizable', 'limited'
    confidence: float = 1.0


@dataclass
class ExternalLink:
    """External reference from Skill 8"""
    id: str
    citation: str
    reference_type: str  # 'supports', 'contradicts', 'extends', 'methodology'
    reference_id: Optional[str] = None
    description: Optional[str] = None
    confidence: float = 1.0


@dataclass
class ContributionNode:
    """Core contribution from Skill 9"""
    id: str
    contribution_type: str  # 'main_finding', 'theoretical', 'methodological', 'practical'
    headline: str
    description: str
    significance: str  # 'So What?' explanation
    confidence: float = 1.0


@dataclass
class FusedKG:
    """Fused knowledge graph from Skill 10"""
    entities: List[Dict]
    relationships: List[Dict]
    conflict_resolutions: List[Dict]  # [{'entity_ids': [], 'resolution': str, 'confidence': float}]
    overall_confidence: float


@dataclass
class QualityReport:
    """Quality audit report from Skill 11"""
    accuracy: float
    reliability: float
    richness: float
    overall_score: float
    orphan_nodes: List[str] = field(default_factory=list)
    issues: List[str] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)


@dataclass
class VisualizationCode:
    """Visualization code from Skill 12"""
    mermaid_code: str
    d3_json: Dict
    cluster_colors: Dict[str, str]


# =============================================================================
# LLM Integration
# =============================================================================

class LLMClient:
    """LLM client wrapper for GPT-4o calls"""
    
    def __init__(self):
        self.model = settings.llm_model
        self.api_key = settings.llm_api_key
        self._client = None
    
    def _get_client(self):
        """Lazy initialization of OpenAI client"""
        if self._client is None:
            try:
                from openai import OpenAI
                self._client = OpenAI(api_key=self.api_key)
            except ImportError:
                logger.warning("OpenAI client not available, using mock responses")
                self._client = None
        return self._client
    
    async def chat(self, messages: List[Dict], temperature: float = 0.3) -> str:
        """
        Send chat completion request to LLM
        
        Args:
            messages: List of message dicts with 'role' and 'content'
            temperature: Sampling temperature (0-1)
        
        Returns:
            LLM response text
        """
        client = self._get_client()
        
        if client is None:
            # Mock response for testing
            return self._mock_response(messages)
        
        try:
            response = client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=temperature
            )
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"LLM call failed: {e}")
            return self._mock_response(messages)
    
    def _mock_response(self, messages: List[Dict]) -> str:
        """Generate mock response for testing"""
        # Return a structured placeholder for testing
        return json.dumps({
            "entities": [],
            "status": "mock_response"
        })
    
    async def extract_json(self, text: str, schema: str) -> Dict:
        """
        Extract structured JSON from text using LLM
        
        Args:
            text: Input text to analyze
            schema: Description of expected JSON structure
        
        Returns:
            Parsed JSON dict
        """
        prompt = f"""Extract structured information from the following text. 
Return ONLY valid JSON matching this schema: {schema}

Text:
{text}

JSON:"""
        
        messages = [
            {"role": "system", "content": "You are a research assistant that extracts structured information from academic text. Always return valid JSON."},
            {"role": "user", "content": prompt}
        ]
        
        response = await self.chat(messages, temperature=0.1)
        
        # Parse JSON from response
        try:
            # Try to find JSON block
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                return json.loads(json_match.group())
            return json.loads(response)
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON: {e}")
            return {"error": str(e), "raw_response": response}


# =============================================================================
# Skill Implementation Classes
# =============================================================================

class Skill1EntityExtractor:
    """Skill 1: Extract core concept entities from text"""
    
    PROMPT_TEMPLATE = """From the following academic text, extract core concept nodes.
- Identify key concepts, terms, and entities (1-3 words each)
- Focus on nouns and noun phrases that represent important concepts
- Return as JSON array with: id, name, entity_type, description, confidence (0-1)

Text:
{text}

Return JSON format:
{{"entities": [{{"id": "E1", "name": "...", "entity_type": "concept", "description": "...", "confidence": 0.9}}]}}"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def extract(self, text: str) -> List[KGEntityNode]:
        """Extract entities from text"""
        prompt = self.PROMPT_TEMPLATE.format(text=text[:8000])  # Limit text length
        
        messages = [
            {"role": "system", "content": "You are a research assistant extracting knowledge graph entities from academic papers."},
            {"role": "user", "content": prompt}
        ]
        
        response = await self.llm.chat(messages, temperature=0.2)
        
        try:
            data = json.loads(response)
            entities = []
            for item in data.get("entities", []):
                entities.append(KGEntityNode(
                    id=item.get("id", f"ENT_{uuid.uuid4().hex[:8].upper()}"),
                    name=item.get("name", ""),
                    entity_type=item.get("entity_type", "concept"),
                    description=item.get("description"),
                    confidence=item.get("confidence", 0.8)
                ))
            return entities
        except json.JSONDecodeError:
            logger.error("Failed to parse entities from LLM response")
            return []


class Skill2TheoryExtractor:
    """Skill 2: Extract theories and frameworks"""
    
    PROMPT_TEMPLATE = """From the following academic text, identify theories, frameworks, and models.
- Identify parent theories and hierarchical relationships
- Note citations and origins
- Return as JSON array with: id, name, theory_type, parent_theory_id, description, source_citation, confidence

Text:
{text}

Return JSON format:
{{"theories": [{{"id": "T1", "name": "...", "theory_type": "parent_theory", "parent_theory_id": null, "description": "...", "source_citation": "...", "confidence": 0.9}}]}}"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def extract(self, text: str) -> List[TheoryNode]:
        """Extract theories from text"""
        prompt = self.PROMPT_TEMPLATE.format(text=text[:8000])
        
        messages = [
            {"role": "system", "content": "You identify theoretical frameworks in academic research."},
            {"role": "user", "content": prompt}
        ]
        
        response = await self.llm.chat(messages, temperature=0.2)
        
        try:
            data = json.loads(response)
            theories = []
            for item in data.get("theories", []):
                theories.append(TheoryNode(
                    id=item.get("id", f"TH_{uuid.uuid4().hex[:8].upper()}"),
                    name=item.get("name", ""),
                    theory_type=item.get("theory_type", "framework"),
                    parent_theory_id=item.get("parent_theory_id"),
                    description=item.get("description"),
                    source_citation=item.get("source_citation"),
                    confidence=item.get("confidence", 0.8)
                ))
            return theories
        except json.JSONDecodeError:
            logger.error("Failed to parse theories from LLM response")
            return []


class Skill3VariableExtractor:
    """Skill 3: Extract research variables (IV, DV, mediator, moderator)"""
    
    PROMPT_TEMPLATE = """From the following academic text, identify research variables and their relationships.
- Identify Independent Variables (IV), Dependent Variables (DV), Mediators, Moderators
- Note directional relationships (positive, negative, U-shape, etc.)
- Return as JSON with variable definitions and relationship types

Text:
{text}

Return JSON format:
{{"variables": [{{"id": "V1", "name": "...", "variable_type": "independent_variable", "description": "...", "relationships": [{{"target": "V2", "type": "positive"}}]}}]}}"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def extract(self, text: str) -> List[Variable]:
        """Extract variables from text"""
        prompt = self.PROMPT_TEMPLATE.format(text=text[:8000])
        
        messages = [
            {"role": "system", "content": "You identify research variables and their statistical relationships."},
            {"role": "user", "content": prompt}
        ]
        
        response = await self.llm.chat(messages, temperature=0.2)
        
        try:
            data = json.loads(response)
            variables = []
            for item in data.get("variables", []):
                var_type_str = item.get("variable_type", "dependent_variable")
                try:
                    var_type = VariableType(var_type_str)
                except ValueError:
                    var_type = VariableType.DV
                
                relationships = []
                for rel in item.get("relationships", []):
                    rel_type_str = rel.get("type", "positive")
                    try:
                        rel_type = RelationshipLabel(rel_type_str)
                    except ValueError:
                        rel_type = RelationshipLabel.POSITIVE
                    relationships.append({
                        'target': rel.get('target', ''),
                        'type': rel_type,
                        'label': rel.get('label', '')
                    })
                
                variables.append(Variable(
                    id=item.get("id", f"VAR_{uuid.uuid4().hex[:8].upper()}"),
                    name=item.get("name", ""),
                    variable_type=var_type,
                    description=item.get("description"),
                    measurement=item.get("measurement"),
                    relationships=relationships,
                    confidence=item.get("confidence", 0.8)
                ))
            return variables
        except json.JSONDecodeError:
            logger.error("Failed to parse variables from LLM response")
            return []


class Skill4ClusterCreator:
    """Skill 4: Create clusters for entities, theories, and variables"""
    
    CLUSTER_TYPES = {
        'context': 'Context/Background information',
        'theory': 'Theoretical framework',
        'methodology': 'Method and measurement',
        'findings': 'Results and findings'
    }
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def create_clusters(
        self, 
        entities: List[KGEntityNode],
        theories: List[TheoryNode],
        variables: List[Variable]
    ) -> ClusterAssignment:
        """Create cluster assignments for all extracted items"""
        
        # Prepare entity data for clustering
        entity_data = []
        for e in entities:
            entity_data.append({"id": e.id, "name": e.name, "type": "entity"})
        for t in theories:
            entity_data.append({"id": t.id, "name": t.name, "type": "theory"})
        for v in variables:
            entity_data.append({"id": v.id, "name": v.name, "type": "variable"})
        
        prompt = f"""Classify the following items into clusters: context, theory, methodology, findings.
- context: Background, setting, definitions
- theory: Theoretical frameworks, models, hypotheses
- methodology: Methods, measurements, procedures
- findings: Results, statistical evidence, conclusions

Items to classify:
{json.dumps(entity_data, indent=2)}

Return JSON:
{{"clusters": {{"context": ["id1", "id2"], "theory": ["id3"], "methodology": ["id4"], "findings": ["id5"]}}}}"""
        
        messages = [
            {"role": "system", "content": "You classify research concepts into semantic clusters."},
            {"role": "user", "content": prompt}
        ]
        
        response = await self.llm.chat(messages, temperature=0.1)
        
        try:
            data = json.loads(response)
            clusters = data.get("clusters", {})
            cluster_labels = {k: self.CLUSTER_TYPES.get(k, k) for k in clusters.keys()}
            
            # Find unassigned
            all_ids = set(e.id for e in entities) | set(t.id for t in theories) | set(v.id for v in variables)
            assigned = set()
            for cluster_ids in clusters.values():
                assigned.update(cluster_ids)
            unassigned = list(all_ids - assigned)
            
            return ClusterAssignment(
                clusters=clusters,
                unassigned=unassigned,
                cluster_labels=cluster_labels
            )
        except json.JSONDecodeError:
            # Fallback: simple clustering
            return self._fallback_clustering(entities, theories, variables)
    
    def _fallback_clustering(self, entities, theories, variables) -> ClusterAssignment:
        """Simple fallback clustering based on entity types"""
        clusters = {
            'theory': [t.id for t in theories],
            'methodology': [v.id for v in variables if v.variable_type in [VariableType.IV, VariableType.CONTROL]],
            'findings': [v.id for v in variables if v.variable_type in [VariableType.DV]],
            'context': [e.id for e in entities]
        }
        cluster_labels = self.CLUSTER_TYPES.copy()
        
        all_ids = set(e.id for e in entities) | set(t.id for t in theories) | set(v.id for v in variables)
        assigned = set()
        for cluster_ids in clusters.values():
            assigned.update(cluster_ids)
        unassigned = list(all_ids - assigned)
        
        return ClusterAssignment(clusters=clusters, unassigned=unassigned, cluster_labels=cluster_labels)


class Skill5ParameterExtractor:
    """Skill 5: Extract research parameters (sample size, tools, location)"""
    
    PARAM_TYPES = ['sample_size', 'location', 'tool', 'method', 'time_period', 'population']
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def extract(self, text: str) -> List[Parameter]:
        """Extract parameters from text"""
        prompt = f"""From the following academic text, extract research parameters:
- Sample size, demographics, location
- Tools, instruments, software used
- Time period of study
- Methodological details

Text:
{text[:6000]}

Return JSON format:
{{"parameters": [{{"id": "P1", "name": "...", "param_type": "sample_size", "value": "500", "unit": "participants", "description": "..."}}]}}"""
        
        messages = [
            {"role": "system", "content": "You extract methodological parameters from academic papers."},
            {"role": "user", "content": prompt}
        ]
        
        response = await self.llm.chat(messages, temperature=0.2)
        
        try:
            data = json.loads(response)
            parameters = []
            for item in data.get("parameters", []):
                parameters.append(Parameter(
                    id=item.get("id", f"PAR_{uuid.uuid4().hex[:8].upper()}"),
                    name=item.get("name", ""),
                    param_type=item.get("param_type", "method"),
                    value=item.get("value", ""),
                    unit=item.get("unit"),
                    description=item.get("description"),
                    confidence=item.get("confidence", 0.8)
                ))
            return parameters
        except json.JSONDecodeError:
            return []


class Skill6EvidenceMapper:
    """Skill 6: Map statistical evidence to variable relationships"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def map_evidence(
        self,
        variables: List[Variable],
        text: str
    ) -> EvidenceMapping:
        """Link statistical results to variable relationships"""
        
        # Extract statistical results from text
        prompt = f"""From the following academic text, extract statistical evidence for hypotheses:
- Identify p-values, betas, effect sizes
- Match evidence to variable relationships (IV->DV, etc.)
- Label hypothesis status: supported, partial, rejected

Variables to match:
{json.dumps([{"id": v.id, "name": v.name, "type": v.variable_type.value} for v in variables])}

Text:
{text[:8000]}

Return JSON format:
{{"evidence_mappings": [{{"source": "V1", "target": "V2", "status": "supported", "p_value": 0.001, "beta": 0.45, "evidence": "..."}}], "overall_finding": "..."}}"""
        
        messages = [
            {"role": "system", "content": "You link statistical evidence to research hypotheses."},
            {"role": "user", "content": prompt}
        ]
        
        response = await self.llm.chat(messages, temperature=0.2)
        
        try:
            data = json.loads(response)
            mappings = []
            for item in data.get("evidence_mappings", []):
                status_str = item.get("status", "supported")
                try:
                    status = HypothesisStatus(status_str)
                except ValueError:
                    status = HypothesisStatus.SUPPORTED
                
                mappings.append({
                    'source': item.get('source', ''),
                    'target': item.get('target', ''),
                    'status': status,
                    'p_value': item.get('p_value'),
                    'beta': item.get('beta'),
                    'effect_size': item.get('effect_size'),
                    'confidence_interval': item.get('confidence_interval'),
                    'evidence': item.get('evidence', '')
                })
            
            return EvidenceMapping(
                variable_relationships=mappings,
                overall_finding=data.get("overall_finding", "")
            )
        except json.JSONDecodeError:
            return EvidenceMapping(variable_relationships=[], overall_finding="")


class Skill7BoundaryExtractor:
    """Skill 7: Extract boundary conditions"""
    
    BOUNDARY_TYPES = ['temporal', 'spatial', 'demographic', 'contextual']
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def extract(self, text: str) -> List[Boundary]:
        """Extract boundary conditions from text"""
        prompt = f"""From the following academic text, identify boundary conditions and limitations:
- Temporal boundaries (time periods, duration)
- Spatial boundaries (locations, settings)
- Demographic boundaries (sample characteristics)
- Contextual limitations

Text:
{text[:6000]}

Return JSON format:
{{"boundaries": [{{"id": "B1", "boundary_type": "temporal", "description": "...", "scope": "generalizable"}}]}}"""
        
        messages = [
            {"role": "system", "content": "You identify research limitations and boundary conditions."},
            {"role": "user", "content": prompt}
        ]
        
        response = await self.llm.chat(messages, temperature=0.2)
        
        try:
            data = json.loads(response)
            boundaries = []
            for item in data.get("boundaries", []):
                boundaries.append(Boundary(
                    id=item.get("id", f"BND_{uuid.uuid4().hex[:8].upper()}"),
                    boundary_type=item.get("boundary_type", "contextual"),
                    description=item.get("description", ""),
                    scope=item.get("scope", "limited"),
                    confidence=item.get("confidence", 0.8)
                ))
            return boundaries
        except json.JSONDecodeError:
            return []


class Skill8ExternalLinkExtractor:
    """Skill 8: Extract external references and citations"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def extract(self, text: str) -> List[ExternalLink]:
        """Extract external references from text"""
        prompt = f"""From the following academic text, identify external citations and references:
- Identify cited works and their relationships
- Note if citations support, contradict, or extend claims
- Include in-text citations (author, year)

Text:
{text[:6000]}

Return JSON format:
{{"external_links": [{{"id": "EX1", "citation": "Author, Year", "reference_type": "supports", "description": "..."}}]}}"""
        
        messages = [
            {"role": "system", "content": "You identify citations and references in academic papers."},
            {"role": "user", "content": prompt}
        ]
        
        response = await self.llm.chat(messages, temperature=0.2)
        
        try:
            data = json.loads(response)
            links = []
            for item in data.get("external_links", []):
                links.append(ExternalLink(
                    id=item.get("id", f"EXT_{uuid.uuid4().hex[:8].upper()}"),
                    citation=item.get("citation", ""),
                    reference_type=item.get("reference_type", "supports"),
                    reference_id=item.get("reference_id"),
                    description=item.get("description"),
                    confidence=item.get("confidence", 0.7)
                ))
            return links
        except json.JSONDecodeError:
            return []


class Skill9ContributionExtractor:
    """Skill 9: Extract core contributions (So What? factor)"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def extract(self, text: str) -> ContributionNode:
        """Extract the main contribution of the paper"""
        prompt = f"""From the following academic text, identify the core contribution:
- What is the main finding or "So What?" factor?
- Is it theoretical, methodological, or practical?
- Why does this matter?

Text:
{text[:6000]}

Return JSON format:
{{"contribution": {{"id": "C1", "contribution_type": "main_finding", "headline": "...", "description": "...", "significance": "..."}}}}"""
        
        messages = [
            {"role": "system", "content": "You identify the key contribution and significance of academic research."},
            {"role": "user", "content": prompt}
        ]
        
        response = await self.llm.chat(messages, temperature=0.2)
        
        try:
            data = json.loads(response)
            contrib = data.get("contribution", {})
            return ContributionNode(
                id=contrib.get("id", f"CTR_{uuid.uuid4().hex[:8].upper()}"),
                contribution_type=contrib.get("contribution_type", "main_finding"),
                headline=contrib.get("headline", ""),
                description=contrib.get("description", ""),
                significance=contrib.get("significance", ""),
                confidence=contrib.get("confidence", 0.8)
            )
        except json.JSONDecodeError:
            return ContributionNode(
                id=f"CTR_{uuid.uuid4().hex[:8].upper()}",
                contribution_type="main_finding",
                headline="",
                description="",
                significance="",
                confidence=0.0
            )


class Skill10KnowledgeFuser:
    """Skill 10: Fuse multiple knowledge graphs"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def fuse(self, kg_list: List[KnowledgeGraph]) -> FusedKG:
        """Merge multiple source knowledge graphs"""
        
        # Collect entities and relationships from all KGs
        all_entities = []
        all_relationships = []
        
        for kg in kg_list:
            if kg.graph_data:
                all_entities.extend(kg.graph_data.get('entities', []))
                all_relationships.extend(kg.graph_data.get('relationships', []))
        
        # Deduplicate entities by name
        entity_map = {}
        for entity in all_entities:
            name = entity.get('name', '').lower()
            if name not in entity_map:
                entity_map[name] = entity
            else:
                # Keep higher confidence version
                if entity.get('confidence', 0) > entity_map[name].get('confidence', 0):
                    entity_map[name] = entity
        
        # Generate merged entity list
        merged_entities = list(entity_map.values())
        
        # Resolve relationship conflicts
        relationship_map = {}
        conflicts = []
        
        for rel in all_relationships:
            key = (rel.get('source'), rel.get('target'))
            if key in relationship_map:
                # Conflict: same relationship exists
                existing = relationship_map[key]
                if rel.get('relation_type') != existing.get('relation_type'):
                    conflicts.append({
                        'entity_ids': list(key),
                        'resolution': 'kept_higher_confidence',
                        'confidence': max(rel.get('confidence', 0), existing.get('confidence', 0))
                    })
                    if rel.get('confidence', 0) > existing.get('confidence', 0):
                        relationship_map[key] = rel
                else:
                    # Same type, merge properties
                    merged_rel = existing.copy()
                    merged_rel['confidence'] = max(rel.get('confidence', 0), existing.get('confidence', 0))
                    relationship_map[key] = merged_rel
            else:
                relationship_map[key] = rel
        
        merged_relationships = list(relationship_map.values())
        
        # Calculate overall confidence
        if merged_entities:
            avg_confidence = sum(e.get('confidence', 0) for e in merged_entities) / len(merged_entities)
        else:
            avg_confidence = 0
        
        return FusedKG(
            entities=merged_entities,
            relationships=merged_relationships,
            conflict_resolutions=conflicts,
            overall_confidence=avg_confidence
        )


class Skill11QualityAuditor:
    """Skill 11: Audit knowledge graph quality"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def audit(self, kg: Dict) -> QualityReport:
        """Audit quality of a knowledge graph"""
        
        entities = kg.get('entities', [])
        relationships = kg.get('relationships', [])
        
        # Calculate metrics
        accuracy = self._calculate_accuracy(entities, relationships)
        reliability = self._calculate_reliability(entities)
        richness = self._calculate_richness(entities, relationships)
        
        # Find orphan nodes
        orphan_nodes = self._find_orphans(entities, relationships)
        
        # Generate issues and recommendations
        issues = []
        recommendations = []
        
        if len(entities) < 5:
            issues.append("Low number of entities extracted")
            recommendations.append("Review text extraction settings")
        
        if orphan_nodes:
            issues.append(f"{len(orphan_nodes)} orphan nodes detected")
            recommendations.append("Add relationships to connect isolated nodes")
        
        if not relationships:
            issues.append("No relationships extracted")
            recommendations.append("Check variable relationship extraction")
        
        overall_score = (accuracy + reliability + richness) / 3
        
        return QualityReport(
            accuracy=accuracy,
            reliability=reliability,
            richness=richness,
            overall_score=overall_score,
            orphan_nodes=orphan_nodes,
            issues=issues,
            recommendations=recommendations
        )
    
    def _calculate_accuracy(self, entities: List, relationships: List) -> float:
        """Calculate accuracy score based on confidence values"""
        if not entities:
            return 0.0
        
        entity_confidences = [e.get('confidence', 0) for e in entities]
        avg_conf = sum(entity_confidences) / len(entity_confidences)
        
        # Factor in relationship consistency
        rel_consistency = 1.0 if relationships else 0.5
        
        return (avg_conf + rel_consistency) / 2
    
    def _calculate_reliability(self, entities: List) -> float:
        """Calculate reliability based on verification status"""
        if not entities:
            return 0.0
        
        verified_count = sum(1 for e in entities if e.get('verification_status') == 'verified')
        return verified_count / len(entities) if entities else 0.0
    
    def _calculate_richness(self, entities: List, relationships: List) -> float:
        """Calculate graph richness (entity/relationship diversity)"""
        if not entities:
            return 0.0
        
        entity_types = set(e.get('entity_type') for e in entities)
        relationship_types = set(r.get('relation_type') for r in relationships)
        
        # More diversity = higher richness
        type_diversity = (len(entity_types) / 10 + len(relationship_types) / 10) / 2
        type_diversity = min(type_diversity, 1.0)
        
        # Connectivity ratio
        connectivity = len(relationships) / len(entities) if entities else 0
        connectivity = min(connectivity, 2.0) / 2
        
        return (type_diversity + connectivity) / 2
    
    def _find_orphans(self, entities: List, relationships: List) -> List[str]:
        """Find orphan nodes (no relationships)"""
        connected_ids = set()
        for rel in relationships:
            connected_ids.add(rel.get('source'))
            connected_ids.add(rel.get('target'))
        
        orphans = [e.get('id') for e in entities if e.get('id') not in connected_ids]
        return orphans


class Skill12VisualizationCompiler:
    """Skill 12: Compile visualization code"""
    
    CLUSTER_COLORS = {
        'context': '#4CAF50',      # Green
        'theory': '#2196F3',       # Blue
        'methodology': '#FF9800',  # Orange
        'findings': '#9C27B0',     # Purple
        'contribution': '#F44336', # Red
    }
    
    def __init__(self, llm_client: LLMClient):
        self.llm = llm_client
    
    async def compile(
        self,
        kg: Dict,
        format: str = 'mermaid'
    ) -> VisualizationCode:
        """Generate visualization code"""
        
        entities = kg.get('entities', [])
        relationships = kg.get('relationships', [])
        clusters = kg.get('clusters', {})
        
        mermaid_code = self._generate_mermaid(entities, relationships, clusters)
        d3_json = self._generate_d3(entities, relationships, clusters)
        
        return VisualizationCode(
            mermaid_code=mermaid_code,
            d3_json=d3_json,
            cluster_colors=self.CLUSTER_COLORS
        )
    
    def _generate_mermaid(
        self,
        entities: List[Dict],
        relationships: List[Dict],
        clusters: Dict
    ) -> str:
        """Generate Mermaid.js diagram code"""
        lines = ["graph TD"]
        
        # Color mapping for clusters
        cluster_style = {}
        
        # Define entity nodes
        for entity in entities:
            node_id = entity.get('id', '').replace('-', '_')
            name = entity.get('name', 'Unknown')[:30]
            entity_type = entity.get('entity_type', 'concept')
            
            # Determine cluster color
            cluster = entity.get('cluster', 'context')
            color = self.CLUSTER_COLORS.get(cluster, '#999999')
            
            # Add node with shape based on type
            if entity_type == 'variable':
                lines.append(f'    {node_id}["{name}"]')
            elif entity_type == 'theory':
                lines.append(f'    {node_id}[("{name}")')
            else:
                lines.append(f'    {node_id}("{name}")')
            
            # Add style for cluster coloring
            if cluster not in cluster_style:
                cluster_style[cluster] = color
        
        # Add relationships
        for rel in relationships:
            source = rel.get('source', '').replace('-', '_')
            target = rel.get('target', '').replace('-', '_')
            rel_type = rel.get('relation_type', 'relates_to')
            
            # Simplify relationship type for display
            if rel_type == 'causes':
                arrow = "-->"
            elif rel_type == 'associated_with':
                arrow = "-.-"
            else:
                arrow = "-->"
            
            lines.append(f'    {source} {arrow} {target} : {rel_type}')
        
        # Add cluster styles
        for cluster, color in cluster_style.items():
            lines.append(f'    classDef {cluster} fill:{color}')
        
        return '\n'.join(lines)
    
    def _generate_d3(
        self,
        entities: List[Dict],
        relationships: List[Dict],
        clusters: Dict
    ) -> Dict:
        """Generate D3.js compatible JSON structure"""
        nodes = []
        links = []
        
        # Convert entities to D3 nodes
        for entity in entities:
            cluster = entity.get('cluster', 'context')
            nodes.append({
                'id': entity.get('id', ''),
                'name': entity.get('name', ''),
                'type': entity.get('entity_type', 'concept'),
                'cluster': cluster,
                'color': self.CLUSTER_COLORS.get(cluster, '#999999'),
                'group': list(self.CLUSTER_COLORS.keys()).index(cluster) if cluster in self.CLUSTER_COLORS else 0
            })
        
        # Convert relationships to D3 links
        for rel in relationships:
            links.append({
                'source': rel.get('source', ''),
                'target': rel.get('target', ''),
                'type': rel.get('relation_type', 'relates_to'),
                'strength': rel.get('confidence', 0.5),
                'properties': rel.get('properties', {})
            })
        
        return {
            'nodes': nodes,
            'links': links,
            'clusterColors': self.CLUSTER_COLORS
        }


# =============================================================================
# Main Generator Class
# =============================================================================

class TextKGGenerator:
    """
    Text Knowledge Graph Generator
    
    Converts academic papers and documents into structured knowledge graphs
    using 12 LLM-based Skills.
    
    Usage:
        generator = TextKGGenerator()
        
        # Generate from text
        kg = await generator.generate(paper_text)
        
        # Generate from document file
        kg = await generator.generate_from_document('/path/to/paper.pdf')
    """
    
    def __init__(self):
        self.llm = LLMClient()
        
        # Initialize skill extractors
        self.skill1 = Skill1EntityExtractor(self.llm)
        self.skill2 = Skill2TheoryExtractor(self.llm)
        self.skill3 = Skill3VariableExtractor(self.llm)
        self.skill4 = Skill4ClusterCreator(self.llm)
        self.skill5 = Skill5ParameterExtractor(self.llm)
        self.skill6 = Skill6EvidenceMapper(self.llm)
        self.skill7 = Skill7BoundaryExtractor(self.llm)
        self.skill8 = Skill8ExternalLinkExtractor(self.llm)
        self.skill9 = Skill9ContributionExtractor(self.llm)
        self.skill10 = Skill10KnowledgeFuser(self.llm)
        self.skill11 = Skill11QualityAuditor(self.llm)
        self.skill12 = Skill12VisualizationCompiler(self.llm)
        
        self.session = get_sync_session()
        logger.info("TextKGGenerator initialized")
    
    async def generate(
        self,
        paper_text: str,
        options: Optional[Dict] = None
    ) -> KnowledgeGraph:
        """
        Generate knowledge graph from paper text
        
        Args:
            paper_text: Raw text content of academic paper
            options: Optional generation options (skills to use, etc.)
        
        Returns:
            KnowledgeGraph object with extracted entities and relationships
        """
        options = options or {}
        logger.info(f"Generating KG from text ({len(paper_text)} chars)")
        
        # Generate unique KG ID
        kg_id = f"KG_{uuid.uuid4().hex[:8].upper()}"
        
        # Execute all 12 skills
        try:
            # Skills 1-3: Extract core components
            entities = await self.skill1.extract(paper_text)
            logger.info(f"Skill 1: Extracted {len(entities)} entities")
            
            theories = await self.skill2.extract(paper_text)
            logger.info(f"Skill 2: Extracted {len(theories)} theories")
            
            variables = await self.skill3.extract(paper_text)
            logger.info(f"Skill 3: Extracted {len(variables)} variables")
            
            # Skill 4: Create clusters
            clusters = await self.skill4.create_clusters(entities, theories, variables)
            logger.info(f"Skill 4: Created clusters: {list(clusters.clusters.keys())}")
            
            # Skill 5-9: Extract additional information
            parameters = await self.skill5.extract(paper_text)
            logger.info(f"Skill 5: Extracted {len(parameters)} parameters")
            
            evidence = await self.skill6.map_evidence(variables, paper_text)
            logger.info(f"Skill 6: Mapped {len(evidence.variable_relationships)} evidence items")
            
            boundaries = await self.skill7.extract(paper_text)
            logger.info(f"Skill 7: Extracted {len(boundaries)} boundaries")
            
            external_links = await self.skill8.extract(paper_text)
            logger.info(f"Skill 8: Extracted {len(external_links)} external links")
            
            contribution = await self.skill9.extract(paper_text)
            logger.info(f"Skill 9: Extracted contribution")
            
            # Build graph data structure
            kg_data = self._build_kg_data(
                kg_id, entities, theories, variables, clusters,
                parameters, evidence, boundaries, external_links, contribution
            )
            
            # Skill 10: Knowledge fusion (for single KG, just validate)
            fused = await self.skill10.fuse([])
            kg_data['fusion_confidence'] = fused.overall_confidence
            
            # Skill 11: Quality audit
            quality = await self.skill11.audit(kg_data)
            logger.info(f"Skill 11: Quality score = {quality.overall_score:.2f}")
            kg_data['quality_report'] = {
                'accuracy': quality.accuracy,
                'reliability': quality.reliability,
                'richness': quality.richness,
                'overall_score': quality.overall_score,
                'orphan_nodes': quality.orphan_nodes,
                'issues': quality.issues,
                'recommendations': quality.recommendations
            }
            
            # Skill 12: Visualization
            viz = await self.skill12.compile(kg_data)
            kg_data['visualization'] = {
                'mermaid': viz.mermaid_code,
                'd3': viz.d3_json
            }
            
            # Create database record
            kg_record = KnowledgeGraph(
                id=kg_id,
                name=f"Paper KG {kg_id}",
                description=f"Knowledge graph generated from text ({len(paper_text)} chars)",
                graph_type='text',
                entity_count=len(entities) + len(theories) + len(variables),
                relationship_count=len(kg_data.get('relationships', [])),
                source_documents=['text_input'],
                quality_score=quality.overall_score,
                last_audit_at=datetime.utcnow(),
                graph_data=kg_data
            )
            
            self.session.add(kg_record)
            self.session.commit()
            
            logger.info(f"Knowledge graph {kg_id} saved to database")
            return kg_record
            
        except Exception as e:
            logger.error(f"Failed to generate KG: {e}")
            self.session.rollback()
            raise
    
    async def generate_from_document(
        self,
        doc_path: Union[str, Path],
        options: Optional[Dict] = None
    ) -> KnowledgeGraph:
        """
        Generate knowledge graph from document file
        
        Args:
            doc_path: Path to document (PDF, DOCX, TXT)
            options: Optional generation options
        
        Returns:
            KnowledgeGraph object
        """
        doc_path = Path(doc_path)
        
        if not doc_path.exists():
            raise FileNotFoundError(f"Document not found: {doc_path}")
        
        logger.info(f"Processing document: {doc_path}")
        
        # Extract text based on file type
        if doc_path.suffix.lower() == '.txt':
            text = doc_path.read_text(encoding='utf-8')
        elif doc_path.suffix.lower() in ['.pdf', '.docx']:
            # For PDF/DOCX, use text extraction
            text = await self._extract_text_from_document(doc_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_path.suffix}")
        
        # Generate KG from extracted text
        kg = await self.generate(text, options)
        
        # Update source document info
        kg.source_documents = [str(doc_path)]
        kg.name = f"Document KG: {doc_path.name}"
        self.session.commit()
        
        return kg
    
    async def _extract_text_from_document(self, doc_path: Path) -> str:
        """Extract text from PDF or DOCX documents"""
        try:
            if doc_path.suffix.lower() == '.pdf':
                import pypdf
                reader = pypdf.PdfReader(doc_path)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text())
                return '\n'.join(text_parts)
            
            elif doc_path.suffix.lower() == '.docx':
                from docx import Document
                doc = Document(doc_path)
                return '\n'.join([p.text for p in doc.paragraphs])
        
        except ImportError as e:
            logger.warning(f"Missing library for document extraction: {e}")
            return f"[Document content from {doc_path.name} - extraction not available]"
        except Exception as e:
            logger.error(f"Failed to extract text: {e}")
            return f"[Document content from {doc_path.name} - extraction failed]"
    
    def _build_kg_data(
        self,
        kg_id: str,
        entities: List[KGEntityNode],
        theories: List[TheoryNode],
        variables: List[Variable],
        clusters: ClusterAssignment,
        parameters: List[Parameter],
        evidence: EvidenceMapping,
        boundaries: List[Boundary],
        external_links: List[ExternalLink],
        contribution: ContributionNode
    ) -> Dict:
        """Build complete KG data structure"""
        
        # Convert entities to dict format
        kg_entities = []
        for e in entities:
            kg_entities.append({
                'id': e.id,
                'name': e.name,
                'entity_type': e.entity_type,
                'description': e.description,
                'cluster': 'context',
                'confidence': e.confidence,
                'source_section': e.source_section
            })
        
        # Add theories
        for t in theories:
            kg_entities.append({
                'id': t.id,
                'name': t.name,
                'entity_type': 'theory',
                'description': t.description,
                'theory_type': t.theory_type,
                'parent_theory_id': t.parent_theory_id,
                'source_citation': t.source_citation,
                'cluster': 'theory',
                'confidence': t.confidence
            })
        
        # Add variables
        for v in variables:
            kg_entities.append({
                'id': v.id,
                'name': v.name,
                'entity_type': 'variable',
                'variable_type': v.variable_type.value,
                'description': v.description,
                'measurement': v.measurement,
                'cluster': 'methodology' if v.variable_type in [VariableType.IV, VariableType.CONTROL] else 'findings',
                'confidence': v.confidence
            })
        
        # Add parameters
        for p in parameters:
            kg_entities.append({
                'id': p.id,
                'name': p.name,
                'entity_type': 'parameter',
                'param_type': p.param_type,
                'value': p.value,
                'unit': p.unit,
                'description': p.description,
                'cluster': 'methodology',
                'confidence': p.confidence
            })
        
        # Add contribution
        kg_entities.append({
            'id': contribution.id,
            'name': contribution.headline,
            'entity_type': 'contribution',
            'contribution_type': contribution.contribution_type,
            'description': contribution.description,
            'significance': contribution.significance,
            'cluster': 'contribution',
            'confidence': contribution.confidence
        })
        
        # Add boundaries as isolated nodes
        for b in boundaries:
            kg_entities.append({
                'id': b.id,
                'name': b.description[:50],
                'entity_type': 'boundary',
                'boundary_type': b.boundary_type,
                'description': b.description,
                'scope': b.scope,
                'cluster': 'context',
                'confidence': b.confidence
            })
        
        # Add external links
        for link in external_links:
            kg_entities.append({
                'id': link.id,
                'name': link.citation,
                'entity_type': 'external_ref',
                'reference_type': link.reference_type,
                'description': link.description,
                'cluster': 'context',
                'confidence': link.confidence
            })
        
        # Build relationships
        kg_relationships = []
        
        # Variable relationships
        for v in variables:
            for rel in v.relationships:
                kg_relationships.append({
                    'source': v.id,
                    'target': rel.get('target', ''),
                    'relation_type': 'affects',
                    'properties': {
                        'relationship_label': rel.get('type', RelationshipLabel.POSITIVE).value,
                        'direction': 'forward'
                    },
                    'confidence': v.confidence
                })
        
        # Theory hierarchy
        for t in theories:
            if t.parent_theory_id:
                kg_relationships.append({
                    'source': t.id,
                    'target': t.parent_theory_id,
                    'relation_type': 'extends',
                    'confidence': t.confidence
                })
        
        # Evidence mappings
        for ev in evidence.variable_relationships:
            kg_relationships.append({
                'source': ev.get('source', ''),
                'target': ev.get('target', ''),
                'relation_type': 'hypothesis',
                'properties': {
                    'status': ev.get('status', HypothesisStatus.SUPPORTED).value,
                    'p_value': ev.get('p_value'),
                    'beta': ev.get('beta'),
                    'effect_size': ev.get('effect_size')
                },
                'hypothesis_status': ev.get('status', HypothesisStatus.SUPPORTED).value,
                'evidence': ev.get('evidence', ''),
                'confidence': 0.9
            })
        
        return {
            'kg_id': kg_id,
            'entities': kg_entities,
            'relationships': kg_relationships,
            'clusters': clusters.clusters,
            'cluster_labels': clusters.cluster_labels,
            'parameters': [{'id': p.id, 'name': p.name, 'param_type': p.param_type, 'value': p.value} for p in parameters],
            'boundaries': [{'id': b.id, 'boundary_type': b.boundary_type, 'description': b.description} for b in boundaries],
            'external_links': [{'id': l.id, 'citation': l.citation, 'reference_type': l.reference_type} for l in external_links],
            'contribution': {
                'id': contribution.id,
                'headline': contribution.headline,
                'significance': contribution.significance
            }
        }
    
    def get_kg_by_id(self, kg_id: str) -> Optional[KnowledgeGraph]:
        """Retrieve a knowledge graph by ID"""
        return self.session.query(KnowledgeGraph).filter_by(id=kg_id).first()
    
    def list_kgs(self, limit: int = 50) -> List[KnowledgeGraph]:
        """List all knowledge graphs"""
        return self.session.query(KnowledgeGraph).limit(limit).all()
    
    def delete_kg(self, kg_id: str) -> bool:
        """Delete a knowledge graph"""
        kg = self.get_kg_by_id(kg_id)
        if kg:
            self.session.delete(kg)
            self.session.commit()
            return True
        return False


# =============================================================================
# Utility Functions
# =============================================================================

def extract_text_from_pdf(pdf_path: Union[str, Path]) -> str:
    """Extract text from PDF file"""
    import pypdf
    reader = pypdf.PdfReader(str(pdf_path))
    return '\n'.join([page.extract_text() for page in reader.pages])


def extract_text_from_docx(docx_path: Union[str, Path]) -> str:
    """Extract text from DOCX file"""
    from docx import Document
    doc = Document(str(docx_path))
    return '\n'.join([p.text for p in doc.paragraphs])


# =============================================================================
# Async Entry Point
# =============================================================================

async def generate_kg_from_text(
    paper_text: str,
    options: Optional[Dict] = None
) -> KnowledgeGraph:
    """Async convenience function to generate KG from text"""
    generator = TextKGGenerator()
    return await generator.generate(paper_text, options)


async def generate_kg_from_document(
    doc_path: Union[str, Path],
    options: Optional[Dict] = None
) -> KnowledgeGraph:
    """Async convenience function to generate KG from document"""
    generator = TextKGGenerator()
    return await generator.generate_from_document(doc_path, options)